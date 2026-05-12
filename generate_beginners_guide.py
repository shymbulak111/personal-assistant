"""Convert BEGINNERS_GUIDE.md to a student-friendly DOCX."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

def sf(run, size=12, bold=False, italic=False, color=None, mono=False):
    run.font.name = 'Courier New' if mono else 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text='', size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent=True, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        _add_inline(p, text, size=size, bold=bold, italic=italic)
    return p

def _add_inline(p, text, size=12, bold=False, italic=False):
    """Add text with **bold** and `code` inline formatting."""
    pat = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            sf(r, size=size, bold=bold, italic=italic)
        if m.group(2):  # **bold**
            r = p.add_run(m.group(2))
            sf(r, size=size, bold=True)
        elif m.group(3):  # `code`
            r = p.add_run(m.group(3))
            sf(r, size=10, mono=True)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        sf(r, size=size, bold=bold, italic=italic)

def add_heading(text, level=1):
    sizes  = {1: 16, 2: 14, 3: 13}
    before = {1: 16, 2: 10, 3: 8}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before.get(level, 8))
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    sf(r, size=sizes.get(level, 12), bold=True)
    return p

def add_code(lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run('\n'.join(lines))
    sf(r, size=9, mono=True)

def add_note(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    sf(r, size=11, italic=True, color=(80, 80, 80))

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(1.25 + level * 0.75)
    _add_inline(p, text, size=12)

def add_numbered(text, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    _add_inline(p, f'{num}. {text}', size=12)

def add_table(header_row, data_rows):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    for ri, row in enumerate(data_rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    doc.add_paragraph()

# ── Title page ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run('Министерство образования и науки Республики Казахстан')
sf(r, size=12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ОБЪЯСНЕНИЕ ПРОЕКТА')
sf(r, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('«Персональный Ассистент»')
sf(r, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Руководство для начинающих')
sf(r, size=14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Выполнил: студент группы ИС-22\nБасханбаев Мирас')
sf(r, size=12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Алматы, 2026')
sf(r, size=12)

doc.add_page_break()

# ── Parse markdown ────────────────────────────────────────────────────────────
guide_path = r'C:\Users\diaskwdk\Desktop\ProJeM\BEGINNERS_GUIDE.md'
with open(guide_path, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
numbered_counter = {}

while i < len(lines):
    raw = lines[i].rstrip('\n')
    stripped = raw.strip()

    # Skip blank lines
    if not stripped:
        i += 1
        continue

    # Horizontal rule
    if stripped == '---':
        i += 1
        continue

    # Code block
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_code(code_lines)
        i += 1
        continue

    # Heading ##
    m = re.match(r'^(#{1,3})\s+(.*)', stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        add_heading(text, level=level)
        i += 1
        continue

    # Blockquote
    if stripped.startswith('>'):
        text = stripped.lstrip('> ').strip()
        add_note(text)
        i += 1
        continue

    # Table — collect all consecutive table lines
    if stripped.startswith('|'):
        tbl_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl_lines.append(lines[i].strip())
            i += 1
        rows = []
        for tl in tbl_lines:
            cells = [c.strip() for c in tl.strip('|').split('|')]
            rows.append(cells)
        if len(rows) >= 2:
            header = rows[0]
            # rows[1] is separator (---|---), skip
            data = rows[2:]
            add_table(header, data)
        continue

    # Bullet
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = stripped[2:].strip()
        add_bullet(text)
        i += 1
        continue

    # Numbered list
    m2 = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m2:
        num = int(m2.group(1))
        text = m2.group(2)
        add_numbered(text, num)
        i += 1
        continue

    # Regular paragraph
    add_para(stripped)
    i += 1

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\diaskwdk\Desktop\ProJeM\BEGINNERS_GUIDE_Baskhanbaev.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
